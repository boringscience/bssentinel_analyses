"""
edit_original_manuscript.py
Modifies the ORIGINAL manuscript .docx in-place, inserting changes in blue
and marking deletions in red strikethrough. Preserves all original formatting.

Output: MIMIC_revision/MIMIC_Early_Prediction_Manuscript_REVISED.docx
"""

import copy
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0, 0, 180)
RED = RGBColor(180, 0, 0)
BLACK = RGBColor(0, 0, 0)

src = Path(__file__).resolve().parent.parent / "MIMIC_revision" / "MIMIC_Early_Prediction_Manuscript.docx"
dst = Path(__file__).resolve().parent.parent / "MIMIC_revision" / "MIMIC_Early_Prediction_Manuscript_REVISED.docx"


def add_run_after(paragraph, text, color=BLUE, bold=False, italic=False, strike=False):
    """Append a new run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.color.rgb = color
    if bold:
        run.font.bold = bold
    if italic:
        run.font.italic = italic
    if strike:
        run.font.strike = True
    return run


def strike_all_runs(paragraph, color=RED):
    """Mark all existing runs as red strikethrough."""
    for run in paragraph.runs:
        run.font.strike = True
        run.font.color.rgb = color


def insert_paragraph_after(paragraph, text, style=None, color=BLUE):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._element.getparent())
    # Copy style from source paragraph if no explicit style, or find style via document
    if style and hasattr(paragraph, 'part') and paragraph.part:
        try:
            new_para.style = paragraph.part.document.styles[style]
        except Exception:
            pass
    run = new_para.add_run(text)
    run.font.color.rgb = color
    return new_para


def replace_in_paragraph(paragraph, old_text, new_text):
    """Replace text within a paragraph, marking old as red strikethrough and new as blue.
    Works by clearing runs and rebuilding with mixed formatting."""
    full = paragraph.text
    if old_text not in full:
        return False

    parts = full.split(old_text, 1)
    # Clear existing runs
    for run in paragraph.runs:
        run.text = ""
    # First run gets the before text
    if paragraph.runs:
        paragraph.runs[0].text = parts[0]
    else:
        r = paragraph.add_run(parts[0])
        r.font.color.rgb = BLACK

    # Add strikethrough old
    r = paragraph.add_run(old_text)
    r.font.color.rgb = RED
    r.font.strike = True

    # Add blue new
    r = paragraph.add_run(new_text)
    r.font.color.rgb = BLUE

    # Add after text
    if len(parts) > 1 and parts[1]:
        r = paragraph.add_run(parts[1])
        r.font.color.rgb = BLACK

    return True


def clear_and_rewrite(paragraph, segments):
    """Clear paragraph and rewrite with segments: list of (text, color, strike, bold)."""
    # Remove all existing runs
    p_elem = paragraph._element
    for r in list(p_elem.findall(qn('w:r'))):
        p_elem.remove(r)

    for text, color, strike, bold in segments:
        run = paragraph.add_run(text)
        run.font.color.rgb = color
        run.font.strike = strike
        run.font.bold = bold


def main():
    doc = Document(str(src))
    paras = doc.paragraphs

    # ── Para 1: Title ──────────────────────────────────────────────────────
    # Index 1: original title
    p = paras[1]
    strike_all_runs(p)
    add_run_after(p, "\n")
    add_run_after(p, "Development and internal validation of a multi-domain electronic health record machine learning model for detecting clinical deterioration in ICU-linked inpatients: a retrospective study using MIMIC-IV", color=BLUE)

    # ── Para 8: Keywords ──────────────────────────────────────────────────
    p = paras[8]
    replace_in_paragraph(p, "informative missingness; ", "domain ablation; ")
    # Also fix NEWS2 -> mNEWS2 in keywords
    # Need to do it on the rebuilt text
    full = p.text
    if "NEWS2" in full and "mNEWS2" not in full:
        # Find the run containing NEWS2 and fix
        for run in p.runs:
            if "NEWS2" in run.text and "mNEWS2" not in run.text:
                run.text = run.text.replace("NEWS2", "mNEWS2")
                break

    # ── Para 15: Abbreviations ────────────────────────────────────────────
    p = paras[15]
    # Add new abbreviations
    add_run_after(p, " BSS: Brier Skill Score; MEWS: Modified Early Warning Score; qSOFA: quick Sequential Organ Failure Assessment;", color=BLUE)
    # Replace NEWS2 definition
    for run in p.runs:
        if "NEWS2: National Early Warning Score 2" in run.text:
            run.text = run.text.replace(
                "NEWS2: National Early Warning Score 2",
                "mNEWS2: modified National Early Warning Score 2 (computed without supplemental oxygen status)"
            )
            run.font.color.rgb = BLUE
            break

    # ── Para 24: Abstract ──────────────────────────────────────────────────
    p = paras[24]
    old_abstract = p.text

    segments = []
    # Split and rebuild the abstract with changes
    segments.append(("Early warning scores for inpatient clinical deterioration typically rely on contemporaneous vital signs and may not incorporate additional information available in electronic health records (EHRs), including laboratory data, medication context, and patterns of test ordering. We developed and internally validated a machine learning model integrating multi-domain EHR features and evaluated its performance relative to ", BLACK, False, False))
    segments.append(("the National Early Warning Score 2 (NEWS2)", RED, True, False))
    segments.append(("a modified National Early Warning Score 2 (mNEWS2, computed without supplemental oxygen status) and other conventional early warning scores", BLUE, False, False))
    segments.append((". In this retrospective cohort study using MIMIC-IV v3.1 (Beth Israel Deaconess Medical Center, 2008\u20132022), a 6-hourly sliding window framework was applied to adult ICU-linked admissions, yielding 1,334,773 monitoring windows from 80,442 admissions. The composite outcome was ICU transfer, invasive mechanical ventilation, or in-hospital mortality within 24 hours (prevalence 1.96%). An XGBoost model trained on 83 features achieved an AUROC of 0.758 ", BLACK, False, False))
    segments.append(("(95% CI 0.754\u20130.762)", RED, True, False))
    segments.append(("(patient-level design-effect corrected 95% CI 0.754\u20130.763)", BLUE, False, False))
    segments.append((", compared with 0.568 for ", BLACK, False, False))
    segments.append(("NEWS2", RED, True, False))
    segments.append(("mNEWS2", BLUE, False, False))
    segments.append((". ", BLACK, False, False))
    segments.append(("Domain-level ablation analysis confirmed that multi-domain integration adds +0.041 AUROC over a vitals-only model, with vital signs contributing the largest single-domain effect. ", BLUE, False, False))
    segments.append(("At a sensitivity of 90%, the model generated 2.51 alerts per monitored patient-day (positive predictive value 2.8%; negative predictive value 99.5%). ", BLACK, False, False))
    segments.append(("Inclusion of binary laboratory missingness indicators was associated with improved model performance, although these features were treated as predictive signals without causal interpretation. ", RED, True, False))
    segments.append(("Decision curve analysis demonstrated consistent net benefit over ", BLACK, False, False))
    segments.append(("NEWS2", RED, True, False))
    segments.append(("mNEWS2", BLUE, False, False))
    segments.append((" across clinically relevant thresholds. These findings reflect model performance within a single-centre, ICU-linked cohort with internal validation; external validation and prospective evaluation are required prior to clinical implementation.", BLACK, False, False))

    clear_and_rewrite(p, segments)

    # ── Para 28: Intro para 3 — add ablation mention ──────────────────────
    p = paras[28]
    # Insert before ". We report the contribution"
    for run in p.runs:
        if ". We report the contribution" in run.text:
            run.text = run.text.replace(
                ". We report the contribution",
                ". We present comprehensive discrimination, calibration, alert-burden, and decision curve analyses"
            )
            break
    add_run_after(p, ", including domain-level ablation to quantify each clinical domain\u2019s incremental contribution", color=BLUE)

    # ── Para 34: Observation Window — add 6h justification ────────────────
    p = paras[34]
    # Replace NEWS2 with mNEWS2
    for run in p.runs:
        if "NEWS2" in run.text and "mNEWS2" not in run.text:
            run.text = run.text.replace("NEWS2", "mNEWS2")

    # Insert new paragraph after para 34 with 6h justification
    new_p = insert_paragraph_after(p,
        "The 6-hourly assessment interval was selected to approximate standard nursing vital sign assessment frequency on general hospital wards (typically every 4\u20138 hours) and to align with common shift handover intervals. This interval balances temporal resolution against feature independence: shorter windows would increase within-admission correlation and reduce informational content per window, while longer windows would sacrifice clinically meaningful temporal granularity. Similar intervals have been used in published ML deterioration models (Escobar et al. 2020; Churpek et al. 2016). We acknowledge that the combination of 6-hourly windows, 24-hour laboratory lookback, and 24-hour prediction horizon creates overlapping feature histories between consecutive windows; sensitivity analyses using alternative window intervals are recommended in future work.",
        style="Normal", color=BLUE)

    # ── Para 42: Calibration — NEWS2 -> mNEWS2 ───────────────────────────
    p = paras[42]
    for run in p.runs:
        if "NEWS2" in run.text and "mNEWS2" not in run.text:
            run.text = run.text.replace("NEWS2", "mNEWS2")

    # ── Para 43: Statistical Analysis — DEFF correction ───────────────────
    p = paras[43]
    old_stat = p.text
    segments = []
    segments.append(("Statistical Analysis: AUROC confidence intervals were ", BLACK, False, False))
    segments.append(("originally ", RED, True, False))
    segments.append(("computed using window-level bootstrap resampling (n = 1,334,773", BLACK, False, False))
    segments.append(("; 2,000 iterations", BLUE, False, False))
    segments.append(("), which underestimates uncertainty by ignoring within-admission correlation. We applied a design-effect correction: DEFF = 1 + (", BLACK, False, False))
    segments.append(("cluster\u2212size", RED, True, False))
    segments.append(("m\u0304", BLUE, False, False))
    segments.append((" \u2212 1) \u00d7 ICC, where ", BLACK, False, False))
    segments.append(("cluster size is the mean number of analysis windows per admission (\u22484.4)", RED, True, False))
    segments.append(("m\u0304 is the mean number of monitoring windows per admission (16.6; median 9, IQR 5\u201317)", BLUE, False, False))
    segments.append((" and ICC = 0.10 (conservative estimate for physiological EHR trajectories", BLACK, False, False))
    segments.append((" within hospital admissions", BLUE, False, False))
    segments.append(("). This yields DEFF = ", BLACK, False, False))
    segments.append(("1.34 and corrected patient-level 95% CI of 0.754\u20130.762 for the XGBoost AUROC, compared with the window-level CI of 0.755\u20130.761. Both are reported", RED, True, False))
    segments.append(("2.56, and the patient-level corrected 95% CI is computed by multiplying the bootstrap standard error by \u221aDEFF. Both window-level and patient-level CIs are reported", BLUE, False, False))
    segments.append(("; patient-level CIs should be considered the primary estimate. AUPRC was computed using the trapezoidal rule. ", BLACK, False, False))
    segments.append(("Brier score was mean squared error between calibrated predictions and binary outcomes.", RED, True, False))
    segments.append((" ", BLACK, False, False))
    segments.append(("Brier score was computed as mean squared error between cross-validated isotonic-calibrated predictions and binary outcomes, ensuring that the metric reflects post-calibration performance. Brier Skill Score (BSS = 1 \u2212 Brier/Brier_ref, where Brier_ref is the score of a naive classifier predicting event prevalence) quantifies improvement over a no-skill reference.", BLUE, False, False))
    segments.append((" DeLong test(15) for AUROC comparison used the Mann-Whitney U kernel variance formulation. Decision curve analysis net benefit: NB(t) = TP/n \u2212 FP/n \u00d7 t/(1\u2212t)(16).", BLACK, False, False))

    clear_and_rewrite(p, segments)

    # ── Insert NEW Methods subsections after para 46 (System Implementation)
    p_anchor = paras[46]

    p_new3 = insert_paragraph_after(p_anchor,
        "Additional Comparators: In addition to mNEWS2, three conventional early warning scores were computed from available features: quick Sequential Organ Failure Assessment (qSOFA: +1 each for systolic BP \u2264 100 mmHg, respiratory rate \u2265 22, and GCS < 15); Modified Early Warning Score (MEWS: composite of SBP, HR, RR, temperature, and GCS subscores); and Shock Index (heart rate / systolic blood pressure).",
        style="Heading 2", color=BLUE)

    p_new2 = insert_paragraph_after(p_anchor,
        "Sensitivity Analyses: To evaluate the robustness of performance to outcome definition, two sensitivity analyses were conducted: (i) excluding all same-day surgical admissions (n = 8,238 admissions), which are most likely to include planned post-operative ICU transfers; and (ii) restricting to emergency admissions only.",
        style="Heading 2", color=BLUE)

    p_new1 = insert_paragraph_after(p_anchor,
        "Domain-Level Ablation Analysis: To quantify the incremental contribution of each clinical domain, we performed a systematic leave-one-domain-out ablation analysis. Features were grouped into six domains: vital signs (28 features), laboratory values (12 features), medications (3 features), NEWS2 subscores (8 features), missingness indicators (20 features), and temporal/demographics (12 features). For each domain, an XGBoost model was trained on all features except those in the excluded domain, using identical hyperparameters and 5-fold GroupKFold cross-validation. Additionally, domain-only models were trained using features from a single domain. All models used pooled out-of-fold AUROC as the primary metric.",
        style="Heading 2", color=BLUE)

    # ── Para 48: Results cohort — NEWS2 -> mNEWS2 ────────────────────────
    p = paras[48]
    for run in p.runs:
        if "NEWS2" in run.text and "mNEWS2" not in run.text:
            run.text = run.text.replace("NEWS2", "mNEWS2")

    # ── Para 51: Discrimination results — REVISED ─────────────────────────
    p = paras[51]
    old_text = p.text
    segments = []
    segments.append(("Multi-domain EHR feature integration improved discrimination over ", BLACK, False, False))
    segments.append(("NEWS2", RED, True, False))
    segments.append(("mNEWS2", BLUE, False, False))
    segments.append((" (AUROC 0.758 vs. 0.568): bssentinel achieved AUROC 0.758 (patient-level ", BLACK, False, False))
    segments.append(("cluster-adjusted 95% CI 0.754\u20130.762; uncorrected window-level CI 0.755\u20130.761) versus 0.708 (95% CI 0.705\u20130.711) for logistic regression and 0.568 for NEWS2", RED, True, False))
    segments.append(("design-effect corrected 95% CI 0.754\u20130.763; uncorrected window-level CI 0.755\u20130.761) versus 0.708 (95% CI 0.703\u20130.713) for logistic regression and 0.568 for mNEWS2", BLUE, False, False))
    segments.append((" (Table 3; Fig. 2). ", BLACK, False, False))
    segments.append(("Additional conventional scores achieved lower discrimination: qSOFA 0.549 (95% CI 0.546\u20130.552), Shock Index 0.539 (95% CI 0.535\u20130.542), and MEWS 0.533 (95% CI 0.530\u20130.537). ", BLUE, False, False))
    segments.append(("The DeLong test indicated statistically significant differences from both ", BLACK, False, False))
    segments.append(("NEWS2", RED, True, False))
    segments.append(("mNEWS2", BLUE, False, False))
    segments.append((" (z = 80.3, p < 10\u207b\u00b9\u2075) and logistic regression (p < 10\u207b\u00b9\u2075).", BLACK, False, False))

    clear_and_rewrite(p, segments)

    # ── Para 62: NPV paragraph — REVISED ──────────────────────────────────
    p = paras[62]
    strike_all_runs(p)
    add_run_after(p, "\n")
    add_run_after(p, "The NPV of 99.5% at this operating point indicates that a prediction below threshold is associated with a low probability of the composite outcome within 24 hours. However, this high NPV is partly attributable to the low event prevalence (1.96%); a naive classifier predicting no events for all windows would achieve NPV = 98.0%. The incremental NPV of 1.5 percentage points over the no-skill baseline, while statistically meaningful at scale, does not independently justify clinical de-escalation decisions. De-escalation of monitoring intensity is a clinical action that requires prospective safety evaluation, including assessment of the consequences of the 0.5% of false-negative predictions, before implementation.", color=BLUE)

    # ── Para 66: Calibration — REVISED ────────────────────────────────────
    p = paras[66]
    strike_all_runs(p)
    add_run_after(p, "\n")
    add_run_after(p, "Cross-validated isotonic regression calibration produced a Brier score of 0.018 for XGBoost and 0.019 for logistic regression, compared with a naive predict-prevalence reference of 0.019 (Brier Skill Score: XGBoost 0.042, logistic regression 0.023). The reliability diagram (Supplementary Fig. S3) shows close agreement between mean predicted probability and observed fraction positive across all decile bins after calibration. The model\u2019s calibrated probabilities range from 0.001 to 0.08 (maximum ~8% predicted risk), reflecting the low base rate. Clinical interfaces should display ordinal risk percentile ranks or locally calibrated stratum labels rather than absolute probability estimates.", color=BLUE)

    # ── Para 67: Feature importance — add ablation caveat ─────────────────
    p = paras[67]
    # Replace last sentence about 10% gain
    for run in p.runs:
        if "Binary laboratory missingness indicators collectively contributed approximately 10% of total model gain." in run.text:
            run.text = run.text.replace(
                "Binary laboratory missingness indicators collectively contributed approximately 10% of total model gain.",
                "Binary laboratory missingness indicators collectively contributed approximately 10% of total XGBoost gain importance."
            )
            break
    add_run_after(p, " However, the domain ablation analysis (Table 5) demonstrates that removing all 20 missingness indicators produces a negligible AUROC change (0.758 \u2192 0.758; \u0394AUROC = 0.000), indicating that the predictive signal attributed to missingness in single-feature importance rankings is fully captured by correlated features in other domains.", color=BLUE)

    # ── Para 68: Missingness interpretation — REVISED ─────────────────────
    p = paras[68]
    strike_all_runs(p)
    add_run_after(p, "\n")
    add_run_after(p, "While individual feature importance rankings are consistent with the informative missingness hypothesis \u2014 that physician test-ordering decisions may encode acuity assessments \u2014 the domain ablation analysis demonstrates that missingness indicators are redundant with other model features (\u0394AUROC = 0.000 when all 20 missingness features are removed). This redundancy likely arises because the MIMIC-IV cohort contains both ICU and general ward admissions: lactate absence partly encodes ward location rather than patient acuity, and this care-setting information is also captured by other features. The prespecified ward-only subgroup falsification test is required to determine whether missingness indicators provide independent predictive value within a homogeneous care setting; until that analysis is completed, this finding is reported as exploratory and should not be interpreted as evidence that test-ordering behaviour contributes independent predictive signal.", color=BLUE)

    # ── Insert NEW Results subsections after para 71 (DCA) ────────────────
    p_dca = paras[71]
    # NEWS2 -> mNEWS2 in DCA
    for run in p_dca.runs:
        if "NEWS2" in run.text and "mNEWS2" not in run.text:
            run.text = run.text.replace("NEWS2", "mNEWS2")

    p_sens = insert_paragraph_after(p_dca,
        "Sensitivity analyses: Excluding same-day surgical admissions (1,227,791 windows from 72,204 admissions) produced AUROC 0.757 (95% CI 0.754\u20130.760), virtually unchanged from the primary analysis. Restricting to emergency admissions only (712,631 windows) yielded AUROC 0.750 (95% CI 0.746\u20130.754). These results indicate that model discrimination is not substantially driven by planned surgical patient-flow patterns.",
        style="Normal", color=BLUE)

    p_ablation = insert_paragraph_after(p_dca,
        "Domain-level ablation: Vital signs are the dominant predictive domain: removing all 28 vital sign features reduced AUROC from 0.758 to 0.708 (\u0394AUROC = \u22120.051; Table 5). Temporal and demographic features were second (\u0394AUROC = \u22120.014), followed by medications and laboratory values (\u0394AUROC \u2248 \u22120.004 each). Missingness indicators and NEWS2 subscores were fully redundant (\u0394AUROC = 0.000). Domain-only models showed vital signs alone achieved AUROC 0.717; multi-domain integration added +0.041 over this baseline.",
        style="Normal", color=BLUE)

    # ── Para 75: Discussion para 1 — corrected CIs, add comparators ──────
    p = paras[75]
    for run in p.runs:
        if "0.754\u20130.762" in run.text:
            run.text = run.text.replace("0.754\u20130.762", "0.754\u20130.763")
            run.font.color.rgb = BLUE
        if "NEWS2" in run.text and "mNEWS2" not in run.text:
            run.text = run.text.replace("NEWS2", "mNEWS2")
    add_run_after(p, " Among conventional early warning scores, qSOFA (AUROC 0.549), MEWS (0.533), and Shock Index (0.539) performed comparably to or below mNEWS2, consistent with published comparisons in ICU-adjacent populations.", color=BLUE)

    # ── Para 76: AUROC caveats — NEWS2 -> mNEWS2 ─────────────────────────
    p = paras[76]
    for run in p.runs:
        if "NEWS2" in run.text and "mNEWS2" not in run.text:
            run.text = run.text.replace("NEWS2", "mNEWS2")

    # ── Para 77: Alert burden discussion — REVISED ────────────────────────
    p = paras[77]
    strike_all_runs(p)
    add_run_after(p, "\n")
    add_run_after(p, "At 2.51 alerts per monitored patient-day at threshold 0.31, the model generates approximately one alert per patient every 10 hours of monitoring. A PPV of 2.8% means that 97.2% of alerts are false positives \u2014 a substantial clinical burden. Importantly, the comparison with ICU physiological monitor alarm rates (187 per patient-day) is inappropriate: monitor alarms and clinical decision support escalation alerts are cognitively and operationally distinct phenomena. The operational sustainability of 2.51 escalation-class alerts per patient-day requires prospective evaluation.", color=BLUE)
    add_run_after(p, "\n")
    add_run_after(p, "Several design strategies can reduce operational alert burden: (i) suppression windows (no re-alert for 6\u201312 hours after initial alert); (ii) consecutive-alert escalation logic; and (iii) event-level rather than window-level alert definitions. At threshold 0.50, alert burden reduces to 1.01 alerts per patient-day at 61.3% sensitivity \u2014 a trade-off that may be more operationally realistic (21).", color=BLUE)

    # ── Para 78: Missingness discussion — REVISED ─────────────────────────
    p = paras[78]
    strike_all_runs(p)
    add_run_after(p, "\n")
    add_run_after(p, "Individual feature importance analysis identified lactate absence as the highest-ranked predictor (5.8% of XGBoost gain). However, the domain-level ablation analysis (Table 5) substantially qualifies this: removing all 20 missingness indicators produces \u0394AUROC = 0.000, demonstrating full redundancy with other domains. This likely reflects the care-setting confound: lactate absence encodes ward vs ICU location, information also captured by temporal features and vital sign patterns. We emphasise that this finding should not be interpreted as evidence that clinician test-ordering behaviour per se drives prediction. The ward-only subgroup SHAP analysis remains the key follow-up.", color=BLUE)

    # ── Insert multidimensional deterioration paragraph after para 78 ─────
    insert_paragraph_after(paras[79],
        "Recent work has highlighted that clinical deterioration is increasingly understood as a multidimensional process extending beyond physiological instability to encompass changes in care demands, surveillance intensity, workflow patterns, and interprofessional coordination (Cesare & Cocchieri, 2024). The current model relies exclusively on structured physiological, laboratory, and medication data. Future work should evaluate whether profession-specific indicators \u2014 including nursing assessment patterns, care coordination signals, and clinical workflow markers \u2014 provide additional predictive value beyond the domains examined here.",
        style="Normal", color=BLUE)

    # ── Para 80: Limitations — add vasopressor caveat ─────────────────────
    p = paras[80]
    add_run_after(p, " ", color=BLACK)
    add_run_after(p, "Additionally, vasopressor use is included as a predictor while the outcome includes ICU transfer. Active vasopressor therapy may constitute partial circularity; however, the domain ablation shows removing all medication features reduces AUROC by only 0.004. The sensitivity analysis excluding surgical admissions showed minimal impact (AUROC 0.757 vs 0.758), but a full MV+mortality-only outcome analysis remains a priority follow-up.", color=BLUE)
    # Fix missingness limitation
    for run in p.runs:
        if "the missingness finding cannot be attributed to informative missingness without ward-only subgroup validation" in run.text:
            run.text = run.text.replace(
                "the missingness finding cannot be attributed to informative missingness without ward-only subgroup validation",
                "the domain ablation analysis demonstrates that missingness indicators are redundant with other model features; whether they carry independent predictive value requires ward-only subgroup validation"
            )
            run.font.color.rgb = BLUE

    # ── Para 81: Conclusions — REVISED ────────────────────────────────────
    p = paras[81]
    strike_all_runs(p)
    add_run_after(p, "\n")
    add_run_after(p, "In this retrospective cohort of ICU-linked inpatients, a multi-domain XGBoost model achieved AUROC 0.758 (95% CI 0.754\u20130.763), outperforming mNEWS2 (0.568), qSOFA (0.549), MEWS (0.533), and Shock Index (0.539). Domain-level ablation confirmed that vital signs are the dominant predictive domain and that multi-domain integration adds +0.041 AUROC over a vitals-only baseline. Missingness indicators were redundant with other model features. At the 90% sensitivity threshold, the model generates 2.51 alerts per monitored patient-day with PPV 2.8%, representing a substantial false-positive burden requiring prospective evaluation. External validation, outcome-restricted sensitivity analysis, and prospective alert-behaviour studies are required before clinical deployment.", color=BLUE)

    # ── Para 82: References — add new ones ────────────────────────────────
    p = paras[82]
    add_run_after(p, "\n\n")
    add_run_after(p, "24. Churpek MM, Yuen TC, Winslow C, et al. Multicenter comparison of machine learning methods and conventional regression for predicting clinical deterioration on the wards. Crit Care Med. 2016;44(2):368\u2013374.", color=BLUE)
    add_run_after(p, "\n\n")
    add_run_after(p, "25. Cesare M, Cocchieri A. Can an increase in nursing care complexity raise the risk of intra-hospital and intensive care unit transfers in children? J Pediatr Nurs. 2024;80:91\u201399. doi:10.1016/j.pedn.2024.11.015", color=BLUE)

    # ── Save ──────────────────────────────────────────────────────────────
    doc.save(str(dst))
    print(f"Revised manuscript saved to {dst}")
    print("  Blue = new/modified text")
    print("  Red strikethrough = deleted text")
    print("  Black = unchanged text")


if __name__ == "__main__":
    main()
